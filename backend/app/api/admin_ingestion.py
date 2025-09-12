from fastapi import APIRouter, File, UploadFile, Form, HTTPException, Query
from fastapi.responses import JSONResponse
from typing import Optional
import json
import logging
from pathlib import Path
from datetime import datetime
from sqlalchemy import text
import traceback
import io
import PyPDF2
from docx import Document as DocxDocument
from ..processing.video_processor import VideoProcessor
from ..processing.video_chunker import VideoChunker, VideoChunk
from ..core.video_migration import migrate_database_for_video_support

logger = logging.getLogger(__name__)

router = APIRouter()

async def _store_video_chunks(doc_id: int, video_chunks: List[VideoChunk], video_data: Dict[str, Any]) -> int:
    """Store video chunks and segments in database"""
    from ..core.database import db
    import openai
    import os
    from ..core.vector_store import vector_store
    
    chunks_stored = 0
    
    try:
        # Store transcript segments first
        with db.engine.connect() as conn:
            segments = video_data['transcript'].get('segments', [])
            for segment in segments:
                conn.execute(
                    text("""
                        INSERT INTO video_segments 
                        (document_id, segment_id, start_time, end_time, text, 
                         tokens, temperature, avg_logprob, compression_ratio, no_speech_prob)
                        VALUES (:document_id, :segment_id, :start_time, :end_time, :text,
                                :tokens, :temperature, :avg_logprob, :compression_ratio, :no_speech_prob)
                    """),
                    {
                        'document_id': doc_id,
                        'segment_id': segment.get('id', 0),
                        'start_time': segment.get('start', 0.0),
                        'end_time': segment.get('end', 0.0),
                        'text': segment.get('text', ''),
                        'tokens': segment.get('tokens', []),
                        'temperature': segment.get('temperature', 0.0),
                        'avg_logprob': segment.get('avg_logprob', 0.0),
                        'compression_ratio': segment.get('compression_ratio', 0.0),
                        'no_speech_prob': segment.get('no_speech_prob', 0.0)
                    }
                )
            conn.commit()
            logger.info(f"Stored {len(segments)} video segments")
        
        # Process and store video chunks with embeddings
        for i, chunk in enumerate(video_chunks):
            try:
                # Create chunk ID
                chunk_id = f"{doc_id}_video_chunk_{i}"
                
                # Generate embedding
                embedding_response = openai.embeddings.create(
                    model="text-embedding-3-large",
                    input=chunk.content,
                    dimensions=1024
                )
                embedding = embedding_response.data[0].embedding
                
                # Prepare metadata for Pinecone
                pinecone_metadata = {
                    'content': chunk.content,
                    'document_id': str(doc_id),
                    'filename': chunk.metadata.get('filename', ''),
                    'content_type': 'video',
                    'start_time': chunk.start_time,
                    'end_time': chunk.end_time,
                    'duration': chunk.end_time - chunk.start_time,
                    'timestamp_display': chunk.metadata.get('timestamp_display', ''),
                    'language': chunk.metadata.get('language', 'en')
                }
                
                # Store in Pinecone
                vector_store.index.upsert(
                    vectors=[(chunk_id, embedding, pinecone_metadata)]
                )
                
                # Store chunk metadata in database
                with db.engine.connect() as conn:
                    conn.execute(
                        text("""
                            INSERT INTO video_chunks 
                            (document_id, chunk_id, content, start_time, end_time, 
                             segment_ids, timestamp_display, duration_display, avg_confidence, metadata)
                            VALUES (:document_id, :chunk_id, :content, :start_time, :end_time,
                                    :segment_ids, :timestamp_display, :duration_display, :avg_confidence, :metadata)
                        """),
                        {
                            'document_id': doc_id,
                            'chunk_id': chunk_id,
                            'content': chunk.content,
                            'start_time': chunk.start_time,
                            'end_time': chunk.end_time,
                            'segment_ids': chunk.segment_ids,
                            'timestamp_display': chunk.metadata.get('timestamp_display', ''),
                            'duration_display': chunk.metadata.get('duration_display', ''),
                            'avg_confidence': chunk.metadata.get('avg_confidence', 0.5),
                            'metadata': json.dumps(chunk.metadata)
                        }
                    )
                    conn.commit()
                
                chunks_stored += 1
                logger.info(f"Successfully stored video chunk {i+1}/{len(video_chunks)}")
                
            except Exception as chunk_error:
                logger.error(f"Failed to store video chunk {i}: {chunk_error}")
                
        logger.info(f"Successfully stored {chunks_stored} video chunks")
        return chunks_stored
        
    except Exception as e:
        logger.error(f"Error storing video chunks: {e}")
        raise

@router.post("/upload-with-config")
async def upload_with_config(
    file: UploadFile = File(...),
    metadata: str = Form(...),
    chunking_config: str = Form(...)
):
    """
    Upload a document with metadata and chunking configuration to Supabase.
    """
    try:
        from ..core.database import db, supabase
        from ..core.vector_store import vector_store
        from ..processing.smart_chunker import SmartChunker
        import openai
        import os
        
        # Parse JSON strings from form data
        metadata_dict = json.loads(metadata)
        config_dict = json.loads(chunking_config)
        
        logger.info(f"Uploading file: {file.filename}")
        logger.info(f"Metadata: {metadata_dict}")
        logger.info(f"Supabase client available: {supabase is not None}")
        
        # Read file content
        content = await file.read()
        logger.info(f"File size: {len(content)} bytes")
        
        # Create unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_ext = Path(file.filename).suffix or '.pdf'
        safe_name = "".join(c for c in metadata_dict['displayName'] if c.isalnum() or c in ('-', '_')).rstrip()
        final_filename = f"{timestamp}_{safe_name}{original_ext}"
        logger.info(f"Final filename: {final_filename}")
        
        # Determine if this is a video file
        video_processor = VideoProcessor()
        is_video = video_processor.is_video_file(file.filename)
        logger.info(f"File type detected: {'video' if is_video else 'document'}")
        
        file_url = None
        
        # Upload to Supabase Storage
        if supabase:
            try:
                logger.info("Uploading to Supabase Storage...")
                
                # Upload directly without checking buckets
                content_type = "application/pdf"
                if is_video:
                    # Set appropriate content type for videos
                    if original_ext.lower() == '.mp4':
                        content_type = "video/mp4"
                    elif original_ext.lower() == '.webm':
                        content_type = "video/webm"
                    elif original_ext.lower() == '.mov':
                        content_type = "video/quicktime"
                    else:
                        content_type = "video/mp4"  # Default
                
                response = supabase.storage.from_('documents').upload(
                    path=final_filename,
                    file=content,
                    file_options={"content-type": content_type, "upsert": "true"}
                )
                logger.info(f"Upload response: {response}")
                
                # Get the public URL
                file_url = supabase.storage.from_('documents').get_public_url(final_filename)
                logger.info(f"File URL: {file_url}")
                
            except Exception as storage_error:
                logger.error(f"Storage upload failed: {storage_error}")
                logger.error(f"Error type: {type(storage_error)}")
                logger.error(f"Error details: {str(storage_error)}")
                raise HTTPException(status_code=500, detail=f"Storage upload failed: {str(storage_error)}")
        else:
            logger.error("Supabase client is None - cannot upload to storage")
            raise HTTPException(status_code=500, detail="Supabase client not initialized. Check SUPABASE_URL and SUPABASE_KEY in .env")
        
        # Ensure video database schema exists
        if is_video:
            migrate_database_for_video_support()
        
        # Process video if needed
        video_data = None
        if is_video:
            try:
                logger.info("Processing video file...")
                video_data = await video_processor.process_video(content, final_filename)
                logger.info(f"Video processed: {video_data['video_info']['duration']:.2f}s duration")
            except Exception as video_error:
                logger.error(f"Video processing failed: {video_error}")
                raise HTTPException(status_code=400, detail=f"Video processing failed: {str(video_error)}")
        
        # Save metadata to database
        doc_id = None
        if db.engine:
            with db.engine.connect() as conn:
                # Prepare base parameters
                params = {
                    'filename': final_filename,
                    'display_name': metadata_dict['displayName'],
                    'document_type': metadata_dict.get('documentType'),
                    'document_source': metadata_dict.get('documentSource'),
                    'human_capability_domain': metadata_dict.get('humanCapabilityDomain', 'hr'),
                    'allow_download': metadata_dict.get('allowDownload', True),
                    'show_in_viewer': metadata_dict.get('showInViewer', True),
                    'author': metadata_dict.get('author') or None,
                    'publication_date': metadata_dict.get('publicationDate') or None,
                    'description': metadata_dict.get('description') or None,
                    'chunking_config': json.dumps(config_dict),
                    'file_url': file_url,
                    'content_type': 'video' if is_video else 'document'
                }
                
                # Add video-specific parameters
                if is_video and video_data:
                    video_info = video_data['video_info']
                    transcript = video_data['transcript']
                    
                    params.update({
                        'duration': video_info.get('duration'),
                        'video_width': video_info.get('width'),
                        'video_height': video_info.get('height'),
                        'transcript_language': transcript.get('language', 'en'),
                        'has_audio': video_info.get('has_audio', True),
                        'video_format': video_info.get('format')
                    })
                    
                    insert_query = """
                        INSERT INTO admin_documents 
                        (filename, display_name, document_type, document_source, 
                         human_capability_domain, allow_download, show_in_viewer,
                         author, publication_date, description, chunking_config, file_url,
                         content_type, duration, video_width, video_height, 
                         transcript_language, has_audio, video_format)
                        VALUES (:filename, :display_name, :document_type, :document_source,
                                :human_capability_domain, :allow_download, :show_in_viewer,
                                :author, :publication_date, :description, :chunking_config, :file_url,
                                :content_type, :duration, :video_width, :video_height,
                                :transcript_language, :has_audio, :video_format)
                        RETURNING id
                    """
                else:
                    insert_query = """
                        INSERT INTO admin_documents 
                        (filename, display_name, document_type, document_source, 
                         human_capability_domain, allow_download, show_in_viewer,
                         author, publication_date, description, chunking_config, file_url, content_type)
                        VALUES (:filename, :display_name, :document_type, :document_source,
                                :human_capability_domain, :allow_download, :show_in_viewer,
                                :author, :publication_date, :description, :chunking_config, :file_url, :content_type)
                        RETURNING id
                    """
                
                result = conn.execute(text(insert_query), params)
                conn.commit()
                doc_id = result.fetchone()[0]
                logger.info(f"Saved document metadata to database with ID: {doc_id}")
        else:
            logger.error("Database engine is not initialized")
            raise HTTPException(status_code=500, detail="Database connection not available")
        
        # Process document/video with appropriate chunking
        chunks_created = 0
        try:
            if is_video:
                logger.info("Starting video transcript chunking...")
                
                # Use video chunker for videos
                video_chunker = VideoChunker(config_dict)
                video_chunks = video_chunker.chunk_video_transcript(
                    video_data['transcript'], 
                    {'filename': final_filename, 'document_id': doc_id}
                )
                
                # Store video chunks and segments in database
                chunks_created = await _store_video_chunks(doc_id, video_chunks, video_data)
                
            else:
                logger.info("Starting document processing with smart chunking...")
                
                # Extract text based on file type with page tracking
                text_content = ""
                page_map = {}  # Maps character position to page number
                
                if original_ext.lower() == '.pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                current_pos = 0
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    # Track which characters belong to which page
                    for i in range(len(page_text)):
                        page_map[current_pos + i] = page_num + 1  # Pages are 1-indexed
                    
                    text_content += page_text + "\n"
                    current_pos += len(page_text) + 1
                    
            elif original_ext.lower() in ['.docx', '.doc']:
                doc = DocxDocument(io.BytesIO(content))
                # For Word docs, estimate pages (roughly 3000 chars per page)
                text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                chars_per_page = 3000
                for i, char_pos in enumerate(range(0, len(text_content), chars_per_page)):
                    for j in range(chars_per_page):
                        if char_pos + j < len(text_content):
                            page_map[char_pos + j] = i + 1
            else:
                logger.warning(f"Unsupported file type for chunking: {original_ext}")
                text_content = ""
            
            if text_content:
                # Use Sandusky Current's simple chunking approach
                chunk_size = 1000  # Fixed like Sandusky
                overlap = 200       # Fixed like Sandusky
                
                def create_chunks(text, chunk_size=1000, overlap=200):
                    """Create chunks using Sandusky Current's approach"""
                    chunks = []
                    start = 0
                    
                    while start < len(text):
                        end = start + chunk_size
                        chunk_text = text[start:end]
                        
                        if chunk_text.strip():
                            chunks.append(chunk_text.strip())
                        
                        start += chunk_size - overlap
                        if start >= len(text):
                            break
                    
                    return chunks
                
                chunk_texts = create_chunks(text_content, chunk_size, overlap)
                
                # Convert to chunk objects
                chunks = []
                for i, chunk_text in enumerate(chunk_texts):
                    chunk = type('Chunk', (), {
                        'content': chunk_text,
                        'section': 'Main Content',
                        'chunk_type': 'text',
                        'metadata': {
                            'filename': final_filename,
                            'display_name': metadata_dict['displayName'],
                            'document_id': str(doc_id),
                            'source': metadata_dict.get('documentSource'),
                            'type': metadata_dict.get('documentType'),
                            'author': metadata_dict.get('author'),
                            'file_url': file_url
                        }
                    })()
                    chunks.append(chunk)
                
                logger.info(f"Created {len(chunks)} chunks")
                
                # Generate embeddings and store in Pinecone
                openai.api_key = os.getenv("OPENAI_API_KEY")
                
                for i, chunk in enumerate(chunks):
                    try:
                        # Determine page number for this chunk
                        chunk_start_pos = text_content.find(chunk.content[:50])  # Find chunk in original text
                        page_number = 1  # Default
                        if chunk_start_pos != -1 and chunk_start_pos in page_map:
                            page_number = page_map[chunk_start_pos]
                        
                        # Generate embedding
                        embedding_response = openai.embeddings.create(
                            model="text-embedding-3-large",
                            input=chunk.content,
                            dimensions=1024
                        )
                        embedding = embedding_response.data[0].embedding
                        
                        # Store in Pinecone with page number
                        chunk_id = f"{doc_id}_chunk_{i}"
                        try:
                            upsert_response = vector_store.index.upsert(
                                vectors=[(
                                    chunk_id,
                                    embedding,
                                    {
                                        "document_id": str(doc_id),
                                        "filename": final_filename,
                                        "title": metadata_dict['displayName'],
                                        "content": chunk.content[:8000],  # Allow much more content for better RAG
                                        "section": chunk.section or "",
                                        "chunk_type": chunk.chunk_type,
                                        "page_number": page_number,  # Add page number
                                        "chunk_index": i,
                                        "file_url": file_url,
                                        "document_source": metadata_dict.get('documentSource'),
                                        "document_type": metadata_dict.get('documentType')
                                    }
                                )]
                            )
                            logger.info(f"Successfully stored chunk {i} with ID {chunk_id}")
                            chunks_created += 1
                        except Exception as upsert_error:
                            logger.error(f"Failed to upsert chunk {i}: {upsert_error}")
                            raise
                    except Exception as chunk_error:
                        logger.error(f"Error processing chunk {i}: {chunk_error}")
                
                logger.info(f"Successfully processed and stored {chunks_created} chunks in vector database")
                
        except Exception as processing_error:
            logger.error(f"Document processing error: {processing_error}")
            logger.error(traceback.format_exc())
            # Don't fail the upload if processing fails
            logger.warning("Document uploaded but processing failed. Document will not be searchable in chat.")
        
        return {
            "status": "success",
            "filename": final_filename,
            "display_name": metadata_dict['displayName'],
            "file_url": file_url,
            "chunks_created": chunks_created
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/reindex-all")
async def reindex_all_documents():
    """
    Re-process all documents in the database and update vector store.
    Clears old vectors and creates new ones with proper file URLs.
    """
    try:
        from ..core.database import db, supabase
        from ..core.vector_store import vector_store
        from ..processing.smart_chunker import SmartChunker
        import openai
        import os
        
        if not db.engine:
            raise HTTPException(status_code=500, detail="Database not connected")
        
        if not supabase:
            raise HTTPException(status_code=500, detail="Supabase not connected")
        
        logger.info("Starting full reindexing of all documents...")
        
        # First, clear all existing vectors
        logger.info("Clearing existing vectors from Pinecone...")
        try:
            # Delete all vectors
            vector_store.index.delete(delete_all=True)
            logger.info("Cleared all vectors from index")
        except Exception as e:
            logger.error(f"Error clearing vectors: {e}")
        
        documents_processed = 0
        total_chunks = 0
        errors = []
        
        with db.engine.connect() as conn:
            # Get all documents with file URLs
            result = conn.execute(
                text("SELECT * FROM admin_documents WHERE file_url IS NOT NULL ORDER BY created_at DESC")
            )
            
            documents = result.fetchall()
            total_documents = len(documents)
            logger.info(f"Found {total_documents} documents to reindex")
            
            for row in documents:
                try:
                    logger.info(f"Processing: {row.display_name}")
                    
                    # Download file from Supabase
                    file_data = supabase.storage.from_('documents').download(row.filename)
                    
                    # Get chunking config
                    chunking_config = json.loads(row.chunking_config) if row.chunking_config else {
                        'chunkSize': 1000,
                        'chunkOverlap': 200,
                        'keepTablesIntact': True,
                        'keepListsIntact': True
                    }
                    
                    # Extract text with page tracking
                    text_content = ""
                    page_map = {}
                    file_ext = Path(row.filename).suffix.lower()
                    
                    if file_ext == '.pdf':
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                        current_pos = 0
                        
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            
                            for i in range(len(page_text)):
                                page_map[current_pos + i] = page_num + 1
                            
                            text_content += page_text + "\n"
                            current_pos += len(page_text) + 1
                            
                    elif file_ext in ['.docx', '.doc']:
                        doc = DocxDocument(io.BytesIO(file_data))
                        text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        # Estimate pages for Word docs
                        chars_per_page = 3000
                        for i, char_pos in enumerate(range(0, len(text_content), chars_per_page)):
                            for j in range(chars_per_page):
                                if char_pos + j < len(text_content):
                                    page_map[char_pos + j] = i + 1
                    else:
                        logger.warning(f"Unsupported file type: {file_ext}")
                        continue
                    
                    if text_content:
                        # Use Sandusky Current's simple chunking approach
                        chunk_size = 1000  # Fixed like Sandusky
                        overlap = 200       # Fixed like Sandusky
                        
                        def create_chunks(text, chunk_size=1000, overlap=200):
                            """Create chunks using Sandusky Current's approach"""
                            chunks = []
                            start = 0
                            
                            while start < len(text):
                                end = start + chunk_size
                                chunk_text = text[start:end]
                                
                                if chunk_text.strip():
                                    chunks.append(chunk_text.strip())
                                
                                start += chunk_size - overlap
                                if start >= len(text):
                                    break
                            
                            return chunks
                        
                        chunk_texts = create_chunks(text_content, chunk_size, overlap)
                        
                        # Convert to chunk objects
                        chunks = []
                        for i, chunk_text in enumerate(chunk_texts):
                            chunk = type('Chunk', (), {
                                'content': chunk_text,
                                'section': 'Main Content',
                                'chunk_type': 'text',
                                'metadata': {
                                    'filename': row.filename,
                                    'display_name': row.display_name,
                                    'document_id': str(row.id),
                                    'source': row.document_source,
                                    'type': row.document_type,
                                    'author': row.author,
                                    'file_url': row.file_url
                                }
                            })()
                            chunks.append(chunk)
                        
                        # Generate embeddings and store
                        openai.api_key = os.getenv("OPENAI_API_KEY")
                        
                        for i, chunk in enumerate(chunks):
                            try:
                                # Determine page number
                                chunk_start_pos = text_content.find(chunk.content[:50])
                                page_number = 1
                                if chunk_start_pos != -1 and chunk_start_pos in page_map:
                                    page_number = page_map[chunk_start_pos]
                                
                                # Generate embedding
                                embedding_response = openai.embeddings.create(
                                    model="text-embedding-3-large",
                                    input=chunk.content,
                                    dimensions=1024
                                )
                                embedding = embedding_response.data[0].embedding
                                
                                # Store in Pinecone
                                chunk_id = f"{row.id}_chunk_{i}"
                                vector_store.index.upsert(
                                    vectors=[(
                                        chunk_id,
                                        embedding,
                                        {
                                            "document_id": str(row.id),
                                            "filename": row.filename,
                                            "title": row.display_name,
                                            "content": chunk.content[:8000],  # Allow much more content for better RAG
                                            "section": chunk.section or "",
                                            "chunk_type": chunk.chunk_type,
                                            "page_number": page_number,
                                            "chunk_index": i,
                                            "file_url": row.file_url,
                                            "document_source": row.document_source,
                                            "document_type": row.document_type
                                        }
                                    )]
                                )
                                total_chunks += 1
                            except Exception as chunk_error:
                                logger.error(f"Error processing chunk {i} for {row.display_name}: {chunk_error}")
                        
                        documents_processed += 1
                        logger.info(f"Successfully reindexed {row.display_name} with {len(chunks)} chunks")
                    
                except Exception as e:
                    error_msg = f"Failed to reindex {row.display_name}: {str(e)}"
                    logger.error(error_msg)
                    errors.append(error_msg)
        
        result = {
            "status": "success",
            "message": f"Reindexing complete",
            "documents_processed": documents_processed,
            "total_documents": total_documents,
            "total_chunks_created": total_chunks,
            "errors": errors
        }
        
        logger.info(f"Reindexing complete: {result}")
        return result
        
    except Exception as e:
        logger.error(f"Reindexing error: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/clear-vectors")
async def clear_all_vectors():
    """
    Clear all vectors from Pinecone index.
    WARNING: This will remove all document embeddings and require reindexing.
    """
    try:
        from ..core.vector_store import vector_store
        
        logger.info("Clearing all vectors from Pinecone...")
        vector_store.index.delete(delete_all=True)
        logger.info("Successfully cleared all vectors")
        
        return {
            "status": "success",
            "message": "All vectors cleared from index. Run reindex-all to rebuild."
        }
    except Exception as e:
        logger.error(f"Error clearing vectors: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/documents")
async def get_documents(
    page: int = Query(1, ge=1),
    limit: int = Query(50, ge=1, le=100),
    search: Optional[str] = None,
    source: Optional[str] = None,
    doc_type: Optional[str] = None
):
    """
    Get list of documents from database with pagination.
    """
    try:
        from ..core.database import db
        
        logger.info(f"Fetching documents - page: {page}, limit: {limit}, search: {search}, source: {source}, doc_type: {doc_type}")
        
        documents = []
        total_count = 0
        
        if db.engine:
            with db.engine.connect() as conn:
                # Get total count
                count_query = "SELECT COUNT(*) as count FROM admin_documents WHERE 1=1"
                count_params = {}
                
                if search:
                    count_query += " AND LOWER(display_name) LIKE LOWER(:search)"
                    count_params['search'] = f"%{search}%"
                if source and source != 'all':
                    count_query += " AND document_source = :source"
                    count_params['source'] = source
                if doc_type and doc_type != 'all':
                    count_query += " AND document_type = :doc_type"
                    count_params['doc_type'] = doc_type
                
                count_result = conn.execute(text(count_query), count_params)
                total_count = count_result.scalar() or 0
                logger.info(f"Total documents found: {total_count}")
                
                # Get paginated results
                query = """
                    SELECT id, filename, display_name, document_type, document_source,
                           human_capability_domain, allow_download, show_in_viewer,
                           author, publication_date, description, created_at, file_url
                    FROM admin_documents
                    WHERE 1=1
                """
                params = {}
                
                if search:
                    query += " AND LOWER(display_name) LIKE LOWER(:search)"
                    params['search'] = f"%{search}%"
                if source and source != 'all':
                    query += " AND document_source = :source"
                    params['source'] = source
                if doc_type and doc_type != 'all':
                    query += " AND document_type = :doc_type"
                    params['doc_type'] = doc_type
                
                query += " ORDER BY created_at DESC LIMIT :limit OFFSET :offset"
                params['limit'] = limit
                params['offset'] = (page - 1) * limit
                
                result = conn.execute(text(query), params)
                
                for row in result:
                    doc = {
                        "id": str(row.id) if row.id else '',
                        "filename": row.filename or '',
                        "displayName": row.display_name or '',
                        "documentType": row.document_type or 'general',
                        "documentSource": row.document_source or 'institute',
                        "humanCapabilityDomain": row.human_capability_domain or 'hr',
                        "allowDownload": row.allow_download if row.allow_download is not None else True,
                        "showInViewer": row.show_in_viewer if row.show_in_viewer is not None else True,
                        "author": row.author or '',
                        "publicationDate": str(row.publication_date) if row.publication_date else '',
                        "description": row.description or '',
                        "uploadDate": str(row.created_at) if row.created_at else '',
                        "fileUrl": row.file_url or ''
                    }
                    documents.append(doc)
        else:
            logger.error("Database engine not available")
            return {"documents": [], "page": page, "limit": limit, "total": 0, "error": "Database not connected"}
        
        logger.info(f"Returning {len(documents)} documents")
        return {
            "documents": documents, 
            "page": page, 
            "limit": limit,
            "total": total_count
        }
        
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        logger.error(traceback.format_exc())
        return {
            "documents": [], 
            "page": page, 
            "limit": limit,
            "total": 0,
            "error": str(e)
        }

@router.delete("/documents/{filename:path}")
async def delete_document(filename: str):
    """
    Delete a document from Supabase storage, database, and vector store.
    """
    try:
        from urllib.parse import unquote
        from ..core.database import db, supabase
        from ..core.vector_store import vector_store
        
        filename = unquote(filename)
        logger.info(f"Deleting document: {filename}")
        
        # Get document ID for vector deletion
        doc_id = None
        if db.engine:
            with db.engine.connect() as conn:
                result = conn.execute(
                    text("SELECT id FROM admin_documents WHERE filename = :filename"),
                    {'filename': filename}
                )
                row = result.first()
                if row:
                    doc_id = str(row.id)
        
        # Delete from Supabase Storage
        if supabase:
            try:
                response = supabase.storage.from_('documents').remove([filename])
                logger.info(f"Deleted from Supabase storage: {response}")
            except Exception as storage_error:
                logger.warning(f"Could not delete from storage: {storage_error}")
        
        # Delete from vector store
        if doc_id and vector_store.index:
            try:
                # Delete all chunks for this document
                vector_store.index.delete(
                    filter={"document_id": {"$eq": doc_id}}
                )
                logger.info(f"Deleted vectors for document {doc_id}")
            except Exception as vector_error:
                logger.warning(f"Could not delete vectors: {vector_error}")
        
        # Delete from database
        if db.engine:
            with db.engine.connect() as conn:
                result = conn.execute(
                    text("DELETE FROM admin_documents WHERE filename = :filename"),
                    {'filename': filename}
                )
                conn.commit()
                logger.info(f"Deleted {result.rowcount} rows from database")
        else:
            raise HTTPException(status_code=500, detail="Database connection not available")
            
        return {"status": "success", "message": "Document deleted"}
            
    except Exception as e:
        logger.error(f"Delete error: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/documents/{filename:path}/permissions")
async def get_document_permissions(filename: str):
    """
    Get document permissions for viewing and downloading.
    """
    try:
        from urllib.parse import unquote
        from ..core.database import db
        
        filename = unquote(filename)
        logger.info(f"Getting permissions for: {filename}")
        
        if db.engine:
            with db.engine.connect() as conn:
                result = conn.execute(
                    text("""
                        SELECT allow_download, show_in_viewer 
                        FROM admin_documents 
                        WHERE filename = :filename
                    """),
                    {'filename': filename}
                )
                row = result.first()
                
                if row:
                    permissions = {
                        "allowDownload": row.allow_download if row.allow_download is not None else True,
                        "showInViewer": row.show_in_viewer if row.show_in_viewer is not None else True
                    }
                    logger.info(f"Found permissions: {permissions}")
                    return permissions
        
        return {"allowDownload": True, "showInViewer": True}
        
    except Exception as e:
        logger.error(f"Error getting permissions: {str(e)}")
        return {"allowDownload": True, "showInViewer": True}

@router.get("/documents/{filename:path}/download")
async def download_document(filename: str):
    """
    Get download URL for a document from Supabase Storage.
    """
    try:
        from urllib.parse import unquote
        from ..core.database import db, supabase
        
        filename = unquote(filename)
        logger.info(f"Getting download URL for: {filename}")
        
        # Check permissions first
        perms = await get_document_permissions(filename)
        if not perms["allowDownload"]:
            raise HTTPException(status_code=403, detail="Download not permitted")
        
        if supabase:
            try:
                # Create signed URL for secure download
                response = supabase.storage.from_('documents').create_signed_url(
                    path=filename,
                    expires_in=3600  # 1 hour
                )
                logger.info(f"Created signed URL response: {response}")
                
                if 'signedURL' in response:
                    return {"url": response['signedURL']}
                elif 'signedUrl' in response:
                    return {"url": response['signedUrl']}
                elif 'data' in response and 'signedUrl' in response['data']:
                    return {"url": response['data']['signedUrl']}
                else:
                    logger.error(f"Unexpected response format: {response}")
                    # Fallback to public URL if signed URL fails
                    public_url = supabase.storage.from_('documents').get_public_url(filename)
                    return {"url": public_url}
                    
            except Exception as e:
                logger.error(f"Failed to create signed URL: {e}")
                # Try public URL as fallback
                try:
                    public_url = supabase.storage.from_('documents').get_public_url(filename)
                    return {"url": public_url}
                except:
                    raise HTTPException(status_code=500, detail=str(e))
        else:
            raise HTTPException(status_code=500, detail="Storage service not available")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))